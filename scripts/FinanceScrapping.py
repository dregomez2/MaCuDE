#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Jul 22 20:01:09 2020

"""

from html.parser import HTMLParser
import requests
from docx import Document
import re
from docx.shared import RGBColor
 
 
class MyHTMLParser(HTMLParser):
    def __init__(self,docname):
        HTMLParser.__init__(self)
        self.docname=docname
        self.docfile = r"%s.doc"%self.docname
        self.doc=Document()
        self.title = False
        self.code = False
        self.text=''
        self.processing =None
        self.codeprocessing =None
        self.picindex = 1
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.110 Safari/537.36'}
        self.timeout = 5

    def handle_starttag(self, tag, attrs):
        if re.match(r"h(\d)", tag):
            self.title = True
        if tag =="p":
            self.processing = tag
        if tag == "code":
            self.code = True
            self.codeprocessing = tag
 
    def handle_data(self, data):
            if self.title == True:
                self.doc.add_heading(data, level=2)
            if self.processing:
                self.text = self.text + data
            if self.code == True:
                p =self.doc.add_paragraph()
                run=p.add_run(data)
 
    def handle_endtag(self, tag):
        self.title = False
        if tag == self.processing:
            self.doc.add_paragraph(self.text)
 
            self.processing = None
            self.text=''
        if tag == self.codeprocessing:
            self.code =False
 
 
headers = {'User-Agent':
               'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.110 Safari/537.36'}
timeout = 4
url="https://catalog.ufl.edu/search/?P=MAN%204301"
html_response = requests.get(url,headers=headers,timeout=timeout)
myHTMLParser = MyHTMLParser("test")
myHTMLParser.feed(html_response.text)
myHTMLParser.doc.save(myHTMLParser.docfile)
